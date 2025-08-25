import streamlit as st
import google.generativeai as genai
import docx
import PyPDF2
import json
import io
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# ---------- Helpers ----------
def extract_text_from_file(f):
    t = f.type
    if t == "application/pdf":
        r = PyPDF2.PdfReader(f)
        return "\n".join([p.extract_text() for p in r.pages if p.extract_text()])
    elif t in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",
               "application/msword"):
        d = docx.Document(f)
        return "\n".join(p.text for p in d.paragraphs)
    elif t.startswith("text/"):
        return f.read().decode("utf-8")
    return None

def call_gemini(key, prompt, model="gemini-1.5-flash"):
    genai.configure(api_key=key)
    llm = genai.GenerativeModel(model)
    resp = llm.generate_content(prompt)
    return resp.text if hasattr(resp, "text") else str(resp)

def build_pdf(text):
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    textobject = c.beginText(40, 750)
    textobject.setFont("Helvetica", 10)
    for line in text.split("\n"):
        textobject.textLine(line)
    c.drawText(textobject)
    c.showPage()
    c.save()
    buf.seek(0)
    return buf.getvalue()

def build_docx(text):
    buf = io.BytesIO()
    d = docx.Document()
    for line in text.split("\n"):
        d.add_paragraph(line)
    d.save(buf)
    buf.seek(0)
    return buf.getvalue()

# ---------- JD prompts ----------
def build_create_prompt(data):
    generic = f"Our client is a leading organisation in the {data['industry']} sector."
    return f"""
Create a professional Job Description using the details below.
Do NOT mention real company names – use generic references only.

Job Title: {data['job_title']}
Department: {data['department']}
Industry: {data['industry']}
Location: {data['location']}
Work Setup: {data['work_setup']}
Must-Have Skills: {data['must_have_skills']}
Total Experience: {data['total_experience']}
Educational Qualification: {data['education']}
Company Info: {generic}

Return the full JD in plain text.
"""

def build_parse_prompt(jd_text):
    return f"""
You are an expert technical recruiter assistant.

Given the following job description (JD), generate output in structured markdown format with the following 3 sections:

---

### 1. ✅ Search Criteria
- Boolean Keyword String  
- Mandatory Skills/Experience  
- Preferred Skills/Experience  

---

### 2. 🧠 10 Screening Questions and Answers  
Categorize into:
- Domain Expertise  
- Product/Tech Depth  
- Cross-functional/Partner Management  
- Fitment & Motivation  
(Provide ideal answers too)

---

### 3. 🗺️ Source Mapping
- Companies in India (Chennai preferred)  
- Relevant job titles  
- LinkedIn Filters (Title, Skills, Location, Experience)

---

JD:
{jd_text}
"""

# ---------- Streamlit ----------
st.set_page_config("Recruiter AI", layout="wide")
st.title("🤖 Recruiter AI — JD Creation & Parsing")

choice = st.radio("Do you already have a Job Description?", ["No (Create one)", "Yes (Parse existing)"])

# =========== CREATE ============
if choice.startswith("No"):
    st.subheader("📝 Fill details to generate a Job Description")
    col1, col2 = st.columns(2)
    job_title   = col1.text_input("Job Title*")
    department  = col2.text_input("Department / Function*")
    industry    = col1.text_input("Industry*")
    location    = col2.text_input("Location*")
    work_setup  = col1.selectbox("Work Setup*", ["Remote","Hybrid","Onsite"])
    must_have   = col2.text_area("Must-Have Skills*")
    exp         = col1.text_input("Total Experience Required*")
    edu         = col2.text_input("Educational Qualification*")
    api_key     = st.text_input("🔑 Enter your Gemini API Key (for JD generation):", type="password")

    if st.button("🚀 Generate JD"):
        if not all([job_title,department,industry,location,work_setup,must_have,exp,edu,api_key]):
            st.error("All fields marked * are required, and API key.")
        else:
            with st.spinner("Generating JD..."):
                data = {
                  "job_title":job_title,"department":department,"industry":industry,
                  "location":location,"work_setup":work_setup,
                  "must_have_skills":must_have,"total_experience":exp,"education":edu
                }
                jd = call_gemini(api_key, build_create_prompt(data))
                st.session_state.generated_jd = jd
                st.success("✅ Job Description created!")

    if "generated_jd" in st.session_state:
        st.markdown("### 📄 Generated JD")
        st.text_area("JD Preview", st.session_state.generated_jd, height=300)

        # Download as PDF or Word
        pdf_bytes = build_pdf(st.session_state.generated_jd)
        docx_bytes = build_docx(st.session_state.generated_jd)

        colA, colB = st.columns(2)
        colA.download_button("⬇️ Download JD as PDF", pdf_bytes, file_name="Generated_JD.pdf")
        colB.download_button("⬇️ Download JD as Word", docx_bytes, file_name="Generated_JD.docx")

        # Parse option
        if st.button("🔍 Parse This JD"):
            st.session_state.to_parse = st.session_state.generated_jd
            st.session_state.api_key = api_key

# =========== PARSE ============
else:
    st.subheader("📤 Upload and Parse an existing JD")
    file = st.file_uploader("Upload PDF/DOCX/TXT", type=["pdf","docx","txt"])
    api_key = st.text_input("🔑 Enter your Gemini API Key (for JD parsing):", type="password")
    if file and api_key and st.button("🔍 Parse JD"):
        text = extract_text_from_file(file)
        st.session_state.to_parse = text
        st.session_state.api_key = api_key

# ========= When user wants to parse =========
if "to_parse" in st.session_state and "api_key" in st.session_state:
    with st.spinner("Parsing JD..."):
        raw = call_gemini(st.session_state.api_key, build_parse_prompt(st.session_state.to_parse))

    if raw:
        st.session_state.parsed_output = raw
        st.success("Parsed ✅")
    else:
        st.error("❌ No output from model.")
        st.stop()

# ====== Display parsed JD and downloads ======
if "parsed_output" in st.session_state:
    st.markdown("### 📊 Parsed JD Output")
    st.text_area("Preview", st.session_state.parsed_output, height=400)

    # Download parsed results
    pdf_bytes = build_pdf(st.session_state.parsed_output)
    docx_bytes = build_docx(st.session_state.parsed_output)

    col1, col2 = st.columns(2)
    col1.download_button("⬇️ Download Parsed JD as PDF", pdf_bytes, file_name="Parsed_JD.pdf")
    col2.download_button("⬇️ Download Parsed JD as Word", docx_bytes, file_name="Parsed_JD.docx")
